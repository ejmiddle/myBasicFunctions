# Note: you need to be using OpenAI Python v0.27.0 for the code below to work
# Transcription using the OpenAI API
# Includes some compressing, as the API has a 25MB limit

import openai
import os
import tiktoken
import math

from pydub import AudioSegment
from pydub.silence import split_on_silence
from pydub.utils import mediainfo
from dotenv import load_dotenv
load_dotenv()

# openai.api_key = os.environ["OPENAI_API_KEY"]

def split_audio_with_overlap(input_file, output_folder, chunk_duration_ms, overlap_duration_ms):
    # Load the audio file
    audio = AudioSegment.from_file(input_file)

    # Calculate the number of chunks with overlap
    num_chunks = math.ceil(len(audio) / (chunk_duration_ms)) 

    for i in range(num_chunks):
        # Calculate the start and end positions for each chunk
        start_time = i * (chunk_duration_ms - overlap_duration_ms)
        end_time = min(start_time + chunk_duration_ms + overlap_duration_ms, len(audio)) 

        # Get the chunk with overlap
        chunk_with_overlap = audio[start_time:end_time]
        print(start_time / (60*1000) ,  end_time / (60*1000) )
        # Save the chunk to the output folder
        output_file = f"{output_folder}/chunk_{i + 1}.mp3"
        chunk_with_overlap.export(output_file, format="mp3")
    return num_chunks

def generate_summary(text: str) -> str:
    """
    Generate a summary of the provided text using OpenAI API
    """
    # Initialize the OpenAI API client
    openai.api_key = os.environ["OPENAI_API_KEY"]

    # Use GPT to generate a summary
    instructions = "Please extract all people mentioned in the text and give a brief description about who they are and what opinion they have on the topic discussed. Please return a bullet list in german language."
    # instructions = "Where do the problematic parts come from?"
    
    num_tokens = num_tokens_from_string (instructions, "cl100k_base")
    print(num_tokens)
    num_tokens = num_tokens_from_string (text, "cl100k_base")
    print(num_tokens)

    
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        # model="gpt-4",
        messages=[
            {"role": "system", "content": instructions},
            {"role": "user", "content": text}
        ],
        temperature=0.1,
        n=1,
        max_tokens=400,
        presence_penalty=0,
        frequency_penalty=0.1,
    )
    # Return the generated summary
    return response.choices[0].message.content.strip()


def num_tokens_from_string(string: str, encoding_name: str) -> int:
    encoding = tiktoken.get_encoding(encoding_name)
    num_tokens = len(encoding.encode(string))
    return num_tokens


def display_mp3_properties(file_path):
    info = mediainfo(file_path)
    
    # Extract and print properties
    print("File:", file_path)
    print("Title:", info.get("title"))
    print("Artist:", info.get("artist"))
    print("Album:", info.get("album"))
    print("Duration (milliseconds):", info.get("duration"))
    print("Channels:", info.get("channels"))
    print("Sample Rate:", info.get("sample_rate"))
    print("Bitrate:", info.get("bit_rate"))
    print("Codec:", info.get("codec_name"))
    
    # Note: The exact properties available might vary depending on the MP3 file and its metadata.

def reduce_mp3_size(input_file, filename, output_path, target_reduction=0.5, format="mp3", silence_thresh=-50, min_silence_len=1000):
    ## target_reduction = bitrate will be reduced to target_reduction * original bitrate

    audio = AudioSegment.from_mp3(input_file)
    
    # # Splitting audio where silence is longer than min_silence_len (in ms)
    # chunks = split_on_silence(audio, 
    #                           min_silence_len=min_silence_len, 
    #                           silence_thresh=silence_thresh)

    # # Concatenate chunks back together
    # processed_audio = AudioSegment.empty()
    # for chunk in chunks:
    #     processed_audio += chunk
    processed_audio = audio

    # Convert the current bitrate from bps to kbps for simplification
    frame_rate = audio.frame_rate  # Gets the frame rate
    frame_width = audio.frame_width  # Gets the frame width in bytes
    channels = audio.channels  # Gets the number of channels

    # For uncompressed audio, you can derive bitrate as:
    current_bitrate = frame_rate * frame_width * channels / 2 # corresponds to media info output approx
    target_bitrate = int(current_bitrate * target_reduction)

    # Set a minimum bitrate to avoid extremely poor quality
    target_bitrate = max(target_bitrate, 32)  # 32 kbps minimum
    print(current_bitrate)
    print(str(target_bitrate) )

    
    # Export the audio with the reduced bitrate
    output_file = os.path.join(output_path, "compressed_" + filename)
    output_folder = output_path
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    processed_audio.export(output_file, format=format, bitrate=str(target_bitrate))
    print("Compressed file written to ... ", output_file) 
    return output_file


def list_files_in_directory(directory_path):
    with os.scandir(directory_path) as entries:
        return [entry.name for entry in entries if entry.is_file()]

def list_mp3_files(directory_path):
    return [os.path.join(directory_path, file) for file in os.listdir(directory_path) if file.endswith('.mp3')]


# ######################################################################################### 
# ####### Main
# ######################################################################################### 

input_path = "downloaded_podcasts/ploetzblog"
output_path = "transcripts"
tmp_path = "transcripts"


# filename = "PB 45 - Quereinsteiger im Bäckerhandwerk - Sebastian Däuwel von den Brotpuristen aus Speyer.mp3"
filename = "PB 2 - Roggenvollkornbrot im Kasten.mp3"
input_file = os.path.join(input_path, filename) 
compressed_path = os.path.join(tmp_path, "compressed")


output_file = reduce_mp3_size(input_file, filename, compressed_path, target_reduction=0.5)
display_mp3_properties(input_file)
display_mp3_properties(output_file)


filepath = "/Users/andreasschmidt/Downloads/wissenschaftsfreiheit-doku-ueber-drohende-cancel-culture-in-wissenschaft-und-forschung.l.mp3"
filepath = "compressed files/" + filename

audio_file = AudioSegment.from_mp3(filepath)
output_folder_path = "compressed files"

# output_folder_path = "tmp"
# chunk_duration_ms = 10 * 60 * 1000
# overlap_duration_ms = 1 * 60000  # 1 minute
# num_chunks = split_audio_with_overlap(filepath, output_folder_path, chunk_duration_ms, overlap_duration_ms)


filelist = list_mp3_files(output_folder_path)
my_infos = ""
dict_list = []
for name_of_part in filelist:
    print(name_of_part)
    # Save the chunk to the output folder
    audio_file= open( name_of_part, "rb")
    transcript = openai.Audio.transcribe("whisper-1", audio_file)
    trans_dict = transcript.to_dict()
    # print('number of chars in dict = ', len(trans_dict['text']))
    # resp = generate_summary(trans_dict['text'])
    # print(resp)
    # my_infos = my_infos + resp
    dict_list.append(trans_dict)


from docx import Document
for idx, data_dict in enumerate(dict_list):
        print(idx, data_dict)
        doc = Document()
        doc.add_heading(f'Document {idx + 1}', level=1)
        
        for key, value in data_dict.items():
            doc.add_paragraph(f'{key}: {value}')
        
        doc_path = os.path.join(output_folder_path, f'document_{idx + 1}.docx')
        doc.save(doc_path)
        print(f"Saved: {doc_path}")

import difflib

text1 = dict_list[0]["text"]
text2 = dict_list[1]["text"]

d = difflib.Differ()
diff = list(d.compare(text1.split("."), text2.split(".")))

for line in diff:
    print(line)

print(my_infos)




